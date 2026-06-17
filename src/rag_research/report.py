"""Render state-of-the-art reports from typed models. Pure functions returning Markdown;
DOCX export is an optional, lazily-imported extra. No global state."""
from __future__ import annotations

import re

from .litreview import AxisRole, Ficha, NoveltyProfile, Tier


def _j(xs: list[str] | None) -> str:
    """Join a non-empty list, else an em-dash."""
    return ", ".join(xs) if xs else "—"


def _confidence(ficha: Ficha) -> str:
    """Confidence from evidence strength: a full local PDF earns more trust than metadata."""
    if ficha.pdf_status == "local":
        return "🟢 alta"
    return "🟡 media"


def _ficha_line(ficha: Ficha) -> str:
    parts: list[str] = [f"**{ficha.title}**"]
    if ficha.year is not None:
        parts.append(f"({ficha.year})")
    if ficha.venue:
        parts.append(f"— {ficha.venue}")
    parts.append(f"· DOI: {ficha.doi or 'n/a'}")
    parts.append(f"· {_j(ficha.modalities)}")
    if ficha.is_multimodal:
        parts.append("· multimodal")
    parts.append(f"· XAI: {_j(ficha.xai_techniques)}")
    parts.append(f"· {_confidence(ficha)}")
    line = " ".join(parts)
    if ficha.relation_to_paper:
        line += f" — *{ficha.relation_to_paper}*"
    return line


def _tier_section(title: str, blurb: str, fichas: list[Ficha]) -> str:
    lines = [f"### {title}", blurb, ""]
    if not fichas:
        lines.append("_(ninguno)_")
    else:
        lines.extend(f"- {_ficha_line(f)}" for f in fichas)
    lines.append("")
    return "\n".join(lines)


_TIER_META = {
    Tier.T1: ("T1 — Competidores directos",
              "Igualan la novedad completa, incluido el diferenciador. Amenaza alta — citar y diferenciar."),
    Tier.T2: ("T2 — Gemelos arquitectónicos",
              "Mismo espacio (dominio+modalidad) sin el diferenciador. Citar para no dejar hueco en el SOTA."),
    Tier.T3: ("T3 — Anclas de técnica",
              "Aportan el diferenciador en otro contexto. Antecedente imprescindible."),
    Tier.T4: ("T4 — Contexto", "Datasets, población, trasfondo."),
}


def render_reporte(profile: NoveltyProfile, fichas: list[Ficha]) -> str:
    """Full REPORTE.md: novelty matrix, tier tables, data-driven actionables, missing PDFs."""
    lines = [f"# Estado del Arte — {profile.paper_ref}", ""]

    lines += ["## 1. Perfil de novedad", "| Eje | Descripción | Rol |", "|---|---|---|"]
    lines += [f"| {a.name} | {a.description} | {a.role.value} |" for a in profile.axes]
    lines.append("")

    lines += ["## 2. Tiers", ""]
    for tier, (title, blurb) in _TIER_META.items():
        lines.append(_tier_section(title, blurb, [f for f in fichas if f.tier == tier]))

    lines += ["## 3. Accionables", ""]
    actions: list[str] = []
    actions += [f"Diferenciar explícitamente frente a {f.title}." for f in fichas if f.tier == Tier.T1]
    t2_titles = [f.title for f in fichas if f.tier == Tier.T2]
    if t2_titles:
        actions.append(f"Citar gemelos no referenciados: {', '.join(t2_titles)}.")
    actions += [f"Posicionar {f.title} como antecedente del diferenciador." for f in fichas if f.tier == Tier.T3]
    actions.append("Verificar cada afirmación numérica del manuscrito con verify_claim_against_corpus.")
    lines += [f"{i}. {a}" for i, a in enumerate(actions, 1)]
    lines.append("")

    no_pdf = [f for f in fichas if f.pdf_status != "local"]
    lines.append("## 4. Sin PDF local")
    lines += [f"- {f.title}" for f in no_pdf] if no_pdf else ["_(ninguno)_"]
    lines.append("")

    return "\n".join(lines)


def render_v2(profile: NoveltyProfile, fichas: list[Ficha]) -> str:
    """Condensed, publication-ready estado_del_arte_v2.md with confidence labels."""
    lines = [
        f"# Estado del Arte v2 — {profile.paper_ref}", "",
        "**Confianza:** 🟢 alta (PDF completo) · 🟡 media (solo metadata)", "",
        "## Novedad real a vender",
    ]
    lines += [f"- {a.description}" for a in profile.axes if a.role == AxisRole.DIFFERENTIATOR]
    lines.append("")

    for tier, heading in (
        (Tier.T1, "## Competidores (T1)"),
        (Tier.T2, "## Gemelos a citar (T2)"),
        (Tier.T3, "## Anclas (T3)"),
    ):
        sel = [f for f in fichas if f.tier == tier]
        lines.append(heading)
        lines += [f"- {_ficha_line(f)}" for f in sel] if sel else ["_(ninguno)_"]
        lines.append("")

    t4 = [f.title for f in fichas if f.tier == Tier.T4]
    lines += ["## Contexto (T4)", ", ".join(t4) if t4 else "_(ninguno)_", ""]
    return "\n".join(lines)


def md_to_docx(markdown_text: str, out_path: str) -> str:
    """Convert Markdown to a .docx via python-docx (lazy import). Minimal block converter:
    headings, bullets, numbered lists, paragraphs; inline markup stripped to plain text."""
    try:
        from docx import Document
    except ImportError as e:  # pragma: no cover
        raise ImportError("md_to_docx needs python-docx: pip install python-docx") from e

    doc = Document()
    for line in markdown_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        raw = re.sub(r"[*_`]", "", stripped)
        if stripped.startswith("### "):
            doc.add_heading(raw[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(raw[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(raw[2:], level=1)
        elif stripped.startswith("- "):
            doc.add_paragraph(raw[2:], style="List Bullet")
        elif re.match(r"^\d+\. ", stripped):
            doc.add_paragraph(re.sub(r"^\d+\.\s+", "", raw), style="List Number")
        else:
            doc.add_paragraph(raw)
    doc.save(out_path)
    return out_path
